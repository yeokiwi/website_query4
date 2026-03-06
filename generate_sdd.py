"""Generate SDD.docx for Website Change Monitor following the sdd.pdf template."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page setup ────────────────────────────────────────────
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)

# ── Styles helpers ────────────────────────────────────────
def set_font(run, name='Times New Roman', size=12, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    set_font(run, size=12, bold=True)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_subheading(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    set_font(run, size=12, bold=True)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_body(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    set_font(run, size=12)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_bullet(doc, text, indent=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    set_font(run, size=12)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    return p

def add_table_with_header(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        set_font(run, size=11, bold=True)
        # Light grey shading
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'D9D9D9')
        tcPr.append(shd)
    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ''
            run = cell.paragraphs[0].add_run(cell_text)
            set_font(run, size=11)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table

# ══════════════════════════════════════════════════════════
#  COVER PAGE
# ══════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n\n\nSOFTWARE DESIGN DESCRIPTION')
set_font(run, size=16, bold=True)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('FOR\n\nWebsite Change Monitor (WCM)\n\nVersion 1.0\n\n6 March 2026')
set_font(run2, size=14)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  AMENDMENT RECORD
# ══════════════════════════════════════════════════════════
add_heading(doc, 'Amendment Record')
headers = ['S/N', 'Description/Reason for Amendment', 'Page(s) Affected / Type of Amendment*', 'Approved by (Name/Sign/Date)']
rows = [['1', 'Initial release', 'All / I', '']]
add_table_with_header(doc, headers, rows, col_widths=[0.5, 2.5, 2.0, 1.5])
add_body(doc, '*There are 3 types of amendment: Insertion (I), Deletion (D) and Modification (M)')

doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  TABLE OF CONTENTS  (manual)
# ══════════════════════════════════════════════════════════
add_heading(doc, 'TABLE OF CONTENTS')
toc_entries = [
    ('1.', 'SCOPE', '1'),
    ('2.', 'REFERENCED DOCUMENTS', '1'),
    ('3.', 'SOFTWARE DESIGN', '1'),
    ('3.1', 'CSCI Overview', '1'),
    ('3.2', 'CSCI Architecture', '2'),
    ('3.2.1', 'Structure Charts / Component Diagram', '2'),
    ('3.2.2', 'Text Description / Sequence of Operations', '3'),
    ('3.3', 'System States and Modes', '5'),
    ('3.4', 'Module Specification / Class Description', '6'),
    ('4.', 'GLOBAL DATA STRUCTURES (CSCI DATA)', '12'),
    ('5.', 'CSCI DATA FILES', '13'),
    ('6.', 'REQUIREMENTS TRACEABILITY', '14'),
    ('7.', 'SPECIAL NOTES', '14'),
]
for num, title, page in toc_entries:
    p = doc.add_paragraph()
    run = p.add_run(f'{num:<8}{title}')
    set_font(run, size=12)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  1. SCOPE
# ══════════════════════════════════════════════════════════
add_heading(doc, '1.   SCOPE')
add_body(doc,
    'This Software Design Description (SDD) establishes the design for the Computer Software '
    'Configuration Item (CSCI) identified as Website Change Monitor (WCM). The WCM is a '
    'responsive single-page web application that monitors websites for recent changes by '
    'employing a Large Language Model (LLM) with autonomous web-browsing capabilities to '
    'analyse sites and produce structured change reports. '
    'The SDD describes the allocation of CSCI requirements to software components, defines '
    'the software architecture, and specifies the interface, data, and processing '
    'characteristics of each software component in the CSCI design.'
)

# ══════════════════════════════════════════════════════════
#  2. REFERENCED DOCUMENTS
# ══════════════════════════════════════════════════════════
add_heading(doc, '2.   REFERENCED DOCUMENTS')
ref_headers = ['Doc No.', 'Title']
ref_rows = [
    ['[1]', 'README.md – Website Change Monitor Setup and Configuration Guide'],
    ['[2]', 'website.md – URL Monitor List File Format Specification'],
    ['[3]', 'render.yaml – Render.com Deployment Blueprint'],
    ['[4]', 'Anthropic API Reference, https://docs.anthropic.com'],
    ['[5]', 'DeepSeek API Reference, https://platform.deepseek.com/docs'],
    ['[6]', 'Brave Search API Documentation, https://api.search.brave.com/app/documentation'],
    ['[7]', 'SerpAPI Documentation, https://serpapi.com/search-api'],
]
add_table_with_header(doc, ref_headers, ref_rows, col_widths=[0.7, 5.8])

# ══════════════════════════════════════════════════════════
#  3. SOFTWARE DESIGN
# ══════════════════════════════════════════════════════════
add_heading(doc, '3.   SOFTWARE DESIGN')

# 3.1 CSCI Overview
add_subheading(doc, '3.1   CSCI Overview')
add_body(doc,
    'The Website Change Monitor (WCM) CSCI is a self-contained web application that '
    'operates as a two-tier client-server system deployed as a single Node.js process. '
    'The CSCI consists of:'
)
add_bullet(doc, 'A React-based Single-Page Application (SPA) front-end compiled and served as static assets from the backend in production, or independently via a Vite dev server during development.')
add_bullet(doc, 'A Node.js + Express REST/SSE backend that orchestrates an agentic reasoning loop, delegates tool calls to web-search and URL-fetch services, and persists output as Markdown report files on the local filesystem.')
add_bullet(doc, 'Two configurable LLM providers (Anthropic Claude and DeepSeek) and two configurable web search providers (Brave Search and SerpAPI).')
add_body(doc,
    'The WCM CSCI interfaces with external APIs (LLM and search providers) over HTTPS and '
    'has no dependency on any other system CSCI or hardware configuration item (HWCI). '
    'In production deployment on Render.com, the CSCI is mounted with a persistent disk at '
    '/data to retain the website list and generated reports across redeployments.'
)

# System architecture diagram (text-based)
add_body(doc, 'High-level architecture:')
p = doc.add_paragraph()
run = p.add_run(
    '┌─────────────────────────────────────────────────────┐\n'
    '│                  USER BROWSER                       │\n'
    '│  ┌───────────────────────────────────────────────┐  │\n'
    '│  │          React SPA (client/src/)              │  │\n'
    '│  │  Login · URL Query · Chat · Batch · Reports   │  │\n'
    '│  └───────────────────┬───────────────────────────┘  │\n'
    '└──────────────────────┼──────────────────────────────┘\n'
    '                  HTTPS / SSE\n'
    '┌─────────────────────────────────────────────────────┐\n'
    '│            Node.js + Express Backend                │\n'
    '│  ┌─────────────┐  ┌──────────────┐  ┌───────────┐  │\n'
    '│  │  Auth       │  │  Agentic     │  │  Reports  │  │\n'
    '│  │  Middleware │  │  Loop        │  │  API      │  │\n'
    '│  └─────────────┘  └──────┬───────┘  └───────────┘  │\n'
    '│                          │                          │\n'
    '│  ┌──────────────┐  ┌─────┴──────────┐              │\n'
    '│  │ llmService   │  │ webSearch /    │              │\n'
    '│  │ (Anthropic / │  │ fetchUrl       │              │\n'
    '│  │  DeepSeek)   │  │ Services       │              │\n'
    '│  └──────┬───────┘  └─────┬──────────┘              │\n'
    '└─────────┼────────────────┼────────────────────────-─┘\n'
    '          │                │\n'
    '   ┌──────┴──────┐  ┌──────┴──────────────┐\n'
    '   │ LLM APIs    │  │ Search / Fetch APIs  │\n'
    '   │ (Anthropic/ │  │ (Brave / SerpAPI /   │\n'
    '   │  DeepSeek)  │  │  Target Websites)    │\n'
    '   └─────────────┘  └──────────────────────┘'
)
p.style = doc.styles['Normal']
run.font.name = 'Courier New'
run.font.size = Pt(8)

# 3.2 CSCI Architecture
add_subheading(doc, '3.2   CSCI Architecture')
add_body(doc,
    'The WCM uses a layered, event-driven architecture built on the following design principles:'
)
add_bullet(doc, 'Stateless REST + streaming SSE API: The backend exposes a REST API for CRUD operations and Server-Sent Events (SSE) for streaming LLM responses and batch progress back to the client in real time.')
add_bullet(doc, 'Agentic loop pattern: The core intelligence is implemented as a while-loop that repeatedly calls the LLM, executes any tool calls the model requests (web_search or fetch_url), and appends results back into the conversation until the LLM signals end_turn.')
add_bullet(doc, 'Provider abstraction: Both the LLM provider and the search provider are selectable at runtime via environment variables. The llmService normalises DeepSeek\'s OpenAI-compatible response format to Anthropic\'s canonical format so the agentic loop remains provider-agnostic.')
add_bullet(doc, 'File-based persistence: Reports are stored as Markdown files with YAML front-matter on the local filesystem; the website URL list is stored in a single plain-text Markdown file (website.md).')
add_bullet(doc, 'In-memory session management: Authentication tokens (UUID v4) are stored in a server-side Set. Sessions are invalidated on server restart, requiring users to re-authenticate.')
add_bullet(doc, 'SSRF / URL validation: The fetchUrlService enforces an allow-list of URL patterns (HTTPS only, no private RFC-1918 or loopback addresses) to prevent SSRF attacks.')
add_bullet(doc, 'Content truncation: Fetched webpage content is truncated to a configurable maximum (default 8 000 characters) to stay within LLM context limits.')

# 3.2.1
add_subheading(doc, '3.2.1   Structure Charts / Component Diagram')
add_body(doc,
    'The diagram below shows the decomposition of the WCM CSCI into software components '
    'and their relationships.'
)
p = doc.add_paragraph()
run = p.add_run(
    'WCM CSCI\n'
    '├── Frontend (client/)\n'
    '│   ├── App  ──────────────────── Root orchestrator\n'
    '│   ├── LoginPage  ─────────────── Authentication UI\n'
    '│   ├── Header  ────────────────── Navigation bar\n'
    '│   ├── UrlQueryBar  ───────────── Single URL input\n'
    '│   ├── MessageList / MessageBubble ─ Chat display\n'
    '│   ├── ChatInput  ─────────────── Free-form chat input\n'
    '│   ├── BatchMonitorPanel  ──────── Batch monitor panel\n'
    '│   │   └── WebsiteListEditor  ─── URL list CRUD\n'
    '│   ├── ReportsPanel  ───────────── Reports viewer\n'
    '│   ├── ChangePasswordModal  ────── Password change UI\n'
    '│   ├── ToolActivityCard  ───────── Tool call status\n'
    '│   ├── TypingIndicator  ────────── Loading indicator\n'
    '│   ├── useChat (hook)  ─────────── Chat state & SSE\n'
    '│   └── utils/\n'
    '│       ├── authFetch  ─────────── Auth-aware fetch\n'
    '│       ├── storage  ───────────── localStorage helpers\n'
    '│       └── formatTime  ────────── Timestamp formatting\n'
    '└── Backend (server/)\n'
    '    ├── index.js  ──────────────── Express server & routes\n'
    '    │   ├── Auth routes  ───────── /api/login|logout|auth/*\n'
    '    │   ├── Chat route  ────────── POST /api/chat (SSE)\n'
    '    │   ├── Batch routes  ──────── /api/batch-monitor/*\n'
    '    │   ├── Report routes  ──────── /api/reports/*\n'
    '    │   └── runAgenticLoop  ────── Agentic reasoning loop\n'
    '    └── services/\n'
    '        ├── llmService  ────────── LLM provider abstraction\n'
    '        │   ├── AnthropicProvider  Anthropic Claude API\n'
    '        │   └── DeepSeekProvider   DeepSeek API (OpenAI-compat)\n'
    '        ├── webSearchService  ───── Web search abstraction\n'
    '        │   ├── searchWithBrave   Brave Search API\n'
    '        │   └── searchWithSerpApi  SerpAPI\n'
    '        └── fetchUrlService  ────── URL fetching & parsing'
)
p.style = doc.styles['Normal']
run.font.name = 'Courier New'
run.font.size = Pt(9)

# 3.2.2
add_subheading(doc, '3.2.2   Text Description / Sequence of Operations')
add_body(doc, 'The following table describes the purpose of each software component:')

comp_headers = ['Component', 'Purpose']
comp_rows = [
    ['index.js', 'Entry point for the Express server. Registers all REST endpoints and SSE streams, enforces JWT-like session authentication via the requireAuth middleware, and hosts the runAgenticLoop function that drives tool-calling iteration.'],
    ['runAgenticLoop', 'Core agentic loop. Repeatedly calls createMessage (LLM), processes text and tool_use blocks, executes tools via executeTool, appends results back to the conversation, and returns when the model produces an end_turn stop reason or when MAX_TOOL_CALLS is exceeded.'],
    ['llmService', 'Provider-agnostic LLM abstraction. Exposes createMessage() and getProviderInfo(). Internally selects Anthropic or DeepSeek provider at startup based on the LLM_PROVIDER env variable. Contains message-format conversion utilities (convertMessagesToOpenAI, normalizeOpenAIResponse) so both providers emit the same canonical Anthropic-style response.'],
    ['webSearchService', 'Executes web searches via Brave Search API or SerpAPI based on the SEARCH_PROVIDER env variable. Returns an array of { title, url, snippet } objects.'],
    ['fetchUrlService', 'Fetches a public HTTPS URL, applies SSRF protection (isBlockedUrl), strips non-content HTML elements using Cheerio, and extracts content as plain text or Markdown via Turndown. Truncates output to FETCH_MAX_CHARS characters.'],
    ['App (React)', 'Root React component. Owns top-level state: authentication status, active panel (home / batch / reports), dark-mode preference, and current date. Renders the appropriate panel based on state.'],
    ['useChat (hook)', 'Custom React hook that encapsulates chat state (messages array) and the SSE streaming logic for /api/chat. Parses event: text, tool_start, tool_end, and error SSE frames and updates the messages state incrementally.'],
    ['LoginPage', 'Boundary component. Presents username/password form, calls POST /api/login, and invokes the onLogin callback with the session token on success.'],
    ['Header', 'Navigation boundary component. Displays application name, LLM provider badge, dark-mode toggle, and navigation buttons (Home, Batch Monitor, Reports, Change Password, Logout).'],
    ['UrlQueryBar', 'Boundary component. Accepts a URL and a date from the user. On submit, constructs a structured change-analysis prompt and delegates to the useChat sendMessage function.'],
    ['BatchMonitorPanel', 'Boundary/control component. Displays the list of monitored URLs (via WebsiteListEditor), triggers POST /api/batch-monitor via SSE, and renders per-URL streaming progress.'],
    ['WebsiteListEditor', 'Control component. Manages the URL list via GET/POST/PUT/DELETE /api/batch-monitor/websites. Provides inline add/remove/edit and bulk-save operations.'],
    ['ReportsPanel', 'Boundary component. Lists all saved reports (GET /api/reports), supports individual report viewing, PDF export (browser print dialog), and deletion.'],
    ['ChangePasswordModal', 'Boundary component. Presents current-password + new-password form and calls POST /api/auth/change-password.'],
    ['MessageList / MessageBubble', 'Presentation components. Render the conversation history. MessageBubble renders Markdown (via react-markdown + remark-gfm + rehype-highlight) for assistant messages and plain text for user messages.'],
    ['ToolActivityCard', 'Presentation component. Shows tool name, input, status (running / done / failed), and a truncated result for each tool invocation in a message.'],
    ['authFetch (util)', 'Wrapper around the native fetch API that injects the Authorization: Bearer <token> header and invokes a global auth-failure callback (triggering logout) on 401 responses.'],
    ['storage (util)', 'Thin localStorage abstraction for persisting chat messages, UI theme preference, and the session token across browser sessions.'],
]
add_table_with_header(doc, comp_headers, comp_rows, col_widths=[1.8, 4.7])

add_body(doc, '')
add_body(doc, 'Key sequence: Single URL query')
add_body(doc, '1. User enters URL + optional date in UrlQueryBar → App.handleUrlQuery builds a structured prompt.')
add_body(doc, '2. useChat.sendMessageSSE POSTs messages array to POST /api/chat with the Authorization header.')
add_body(doc, '3. Backend runAgenticLoop calls createMessage → LLM responds with tool_use blocks.')
add_body(doc, '4. executeTool invokes webSearch or fetchUrl, appends tool_result to conversation.')
add_body(doc, '5. Loop continues until LLM emits end_turn; each text chunk is streamed as event: text SSE.')
add_body(doc, '6. Frontend accumulates text in the assistant message bubble in real time.')

add_body(doc, '')
add_body(doc, 'Key sequence: Batch monitor run')
add_body(doc, '1. User clicks Run Batch in BatchMonitorPanel → POST /api/batch-monitor.')
add_body(doc, '2. Backend reads website.md, emits batch_start SSE, then iterates over each URL.')
add_body(doc, '3. For each URL, a fresh runAgenticLoop is started; progress is streamed as batch_item_* SSE events.')
add_body(doc, '4. Completed analysis is written to a timestamped Markdown file in the REPORTS_DIR.')
add_body(doc, '5. After all URLs, batch_done SSE is emitted with success/failure counts.')

# 3.3 System States and Modes
add_subheading(doc, '3.3   System States and Modes')
add_body(doc,
    'The WCM operates in the following states. State transitions are driven by user actions '
    'and authentication events.'
)
state_headers = ['State', 'Description', 'Active Components']
state_rows = [
    ['INIT / Auth-check', 'Application loads; stored session token is validated against GET /api/auth/check. A blank screen is displayed while the check is pending.', 'App, storage (loadToken), authFetch'],
    ['UNAUTHENTICATED', 'No valid session token. Login page is displayed. All /api/* endpoints (except /login and /health) return 401.', 'LoginPage, App'],
    ['HOME (Authenticated)', 'Default authenticated state. UrlQueryBar, MessageList, and ChatInput are visible. Single URL queries and free-form chat are available.', 'App, Header, UrlQueryBar, MessageList, MessageBubble, ChatInput, useChat, ToolActivityCard, TypingIndicator'],
    ['BATCH MONITOR (Authenticated)', 'BatchMonitorPanel overlays the home view. User can edit the URL list and trigger a batch monitoring run.', 'App, Header, BatchMonitorPanel, WebsiteListEditor'],
    ['REPORTS (Authenticated)', 'ReportsPanel overlays the home view. User can browse, read, export, and delete reports.', 'App, Header, ReportsPanel'],
    ['CHANGE PASSWORD (Modal)', 'ChangePasswordModal is rendered over the current panel. Does not affect the underlying panel state.', 'ChangePasswordModal'],
]
add_table_with_header(doc, state_headers, state_rows, col_widths=[1.6, 2.7, 2.2])

add_body(doc, '')
add_body(doc, 'Component operation across states:')
mode_headers = ['Component', 'INIT', 'UNAUTH', 'HOME', 'BATCH', 'REPORTS', 'CHANGE PWD']
mode_rows = [
    ['App', 'X', 'X', 'X', 'X', 'X', 'X'],
    ['LoginPage', '', 'X', '', '', '', ''],
    ['Header', '', '', 'X', 'X', 'X', 'X'],
    ['UrlQueryBar', '', '', 'X', '', '', ''],
    ['MessageList / Bubble', '', '', 'X', '', '', ''],
    ['ChatInput', '', '', 'X', '', '', ''],
    ['useChat', '', '', 'X', '', '', ''],
    ['BatchMonitorPanel', '', '', '', 'X', '', ''],
    ['WebsiteListEditor', '', '', '', 'X', '', ''],
    ['ReportsPanel', '', '', '', '', 'X', ''],
    ['ChangePasswordModal', '', '', '', '', '', 'X'],
    ['ToolActivityCard', '', '', 'X', 'X', '', ''],
    ['TypingIndicator', '', '', 'X', 'X', '', ''],
]
add_table_with_header(doc, mode_headers, mode_rows, col_widths=[2.2, 0.5, 0.65, 0.55, 0.6, 0.7, 0.9])
add_body(doc, 'Note: X means the component is active/rendered in that state.')

# 3.4 Module Specification / Class Description
add_subheading(doc, '3.4   Module Specification / Class Description')

def module_spec(doc, name, class_type, responsibilities, attributes, operations):
    add_body(doc, f'─── {name} ───')
    p = doc.add_paragraph()
    r = p.add_run('Class Name: '); set_font(r, bold=True)
    r2 = p.add_run(name); set_font(r2)
    p = doc.add_paragraph()
    r = p.add_run('Class Type: '); set_font(r, bold=True)
    r2 = p.add_run(class_type); set_font(r2)
    p = doc.add_paragraph()
    r = p.add_run('Responsibilities: '); set_font(r, bold=True)
    r2 = p.add_run(responsibilities); set_font(r2)
    if attributes:
        p = doc.add_paragraph()
        r = p.add_run('Attributes: '); set_font(r, bold=True)
        r2 = p.add_run(attributes); set_font(r2)
    p = doc.add_paragraph()
    r = p.add_run('Operations: '); set_font(r, bold=True)
    r2 = p.add_run(operations); set_font(r2)
    doc.add_paragraph()

module_spec(doc,
    name='index.js (ExpressServer)',
    class_type='Control – orchestrates all HTTP request handling and the agentic reasoning loop.',
    responsibilities=(
        'Initialises the Express application; registers authentication, chat, batch-monitor, '
        'and report REST/SSE endpoints; enforces session-based access control; and hosts the '
        'core agentic reasoning loop.'
    ),
    attributes=(
        'activeSessions: Set<string> – in-memory set of valid session tokens.\n'
        'currentPassword: string – runtime-mutable password stored in memory.\n'
        'WEBSITE_FILE: string – absolute path to website.md.\n'
        'REPORTS_DIR: string – absolute path to the reports directory.\n'
        'MAX_TOOL_CALLS: number – maximum LLM tool calls per agentic turn (default 10).'
    ),
    operations=(
        'runAgenticLoop(messages, onText, onToolStart, onToolEnd): Promise<string> – '
        'iteratively calls the LLM, executes tool calls, and streams events until end_turn.\n'
        'executeTool(name, input): Promise<string> – dispatches to webSearch or fetchUrl.\n'
        'getPromptForUrl(url, currentDate): string – builds the standard change-analysis prompt.\n'
        'generateReportFilename(url): string – creates a deduplicated timestamped filename.\n'
        'writeReport(filename, url, content, status): void – persists a report as Markdown.\n'
        'validateUrl(urlString): string|null – validates HTTPS and non-private URL.\n'
        'readWebsiteFile(): string[] – parses website.md into a URL array.\n'
        'writeWebsiteFile(urls): void – serialises a URL array back to website.md.\n'
        'sseWrite(res, event, data): void – emits a single SSE frame.'
    )
)

module_spec(doc,
    name='llmService',
    class_type='Control – LLM provider abstraction.',
    responsibilities=(
        'Selects and initialises the configured LLM provider at startup. Provides a uniform '
        'createMessage() interface regardless of the underlying provider. Converts between '
        'Anthropic and OpenAI message formats as required.'
    ),
    attributes=(
        'LLM_PROVIDER: string – "anthropic" (default) or "deepseek".\n'
        'MAX_TOKENS: number – maximum tokens for LLM response (default 4096).\n'
        'TEMPERATURE: number – LLM sampling temperature (default 0.7).\n'
        'SYSTEM_PROMPT: string – system instruction for the LLM.\n'
        'anthropicTools: ToolDefinition[] – tool schemas in Anthropic format.\n'
        'openaiTools: ToolDefinition[] – tool schemas in OpenAI function-call format.'
    ),
    operations=(
        'createMessage(messages): Promise<AnthropicResponse> – calls the active provider and '
        'returns a normalised Anthropic-format response.\n'
        'getProviderInfo(): {provider, model} – returns active provider metadata.\n'
        'convertMessagesToOpenAI(messages): OpenAIMessage[] – converts Anthropic-format '
        'message history to OpenAI format for the DeepSeek provider.\n'
        'normalizeOpenAIResponse(response): AnthropicResponse – normalises an OpenAI chat '
        'completion to the Anthropic canonical response structure.'
    )
)

module_spec(doc,
    name='webSearchService',
    class_type='Entity / Service – web search integration.',
    responsibilities=(
        'Executes keyword-based web searches using the configured search provider and '
        'returns structured result objects. Enforces result-count limits.'
    ),
    attributes=(
        'SEARCH_PROVIDER: string – "brave" (default) or "serpapi".'
    ),
    operations=(
        'webSearch(query, numResults?): Promise<SearchResult[]> – dispatches to the '
        'configured search provider.\n'
        'searchWithBrave(query, numResults): Promise<SearchResult[]> – queries Brave Search '
        'API and maps results to { title, url, snippet }.\n'
        'searchWithSerpApi(query, numResults): Promise<SearchResult[]> – queries SerpAPI '
        'and maps results to { title, url, snippet }.'
    )
)

module_spec(doc,
    name='fetchUrlService',
    class_type='Entity / Service – SSRF-protected URL fetcher. Contains a critical security algorithm (SSRF prevention).',
    responsibilities=(
        'Fetches the readable content of a public HTTPS URL. Applies SSRF protection by '
        'rejecting private network addresses. Parses HTML using Cheerio, removes '
        'boilerplate elements, and extracts content as plain text or Markdown.'
    ),
    attributes=(
        'FETCH_MAX_CHARS: number – maximum characters of extracted content (default 8000).\n'
        'BLOCKED_PATTERNS: RegExp[] – patterns matching localhost, RFC-1918, link-local, '
        'and IPv6 loopback/private addresses.'
    ),
    operations=(
        'fetchUrl(url, extractMode?): Promise<string> – fetches and extracts webpage content.\n'
        'isBlockedUrl(urlString): boolean – returns true if the URL is non-HTTPS or '
        'matches any blocked address pattern.'
    )
)

module_spec(doc,
    name='App (React Component)',
    class_type='Control – root application controller.',
    responsibilities=(
        'Manages global application state including authentication status, active panel, '
        'dark-mode preference, current date, and LLM provider info. Renders the '
        'appropriate panel (Login / Home / Batch / Reports) based on state.'
    ),
    attributes=(
        'authenticated: boolean – whether the user has a valid session.\n'
        'authChecked: boolean – whether the initial auth check has completed.\n'
        'darkMode: boolean – current theme preference.\n'
        'showBatch: boolean – whether the Batch Monitor panel is active.\n'
        'showReports: boolean – whether the Reports panel is active.\n'
        'showChangePassword: boolean – whether the Change Password modal is open.\n'
        'batchRunning: boolean – whether a batch monitor run is in progress.\n'
        'providerInfo: {provider, model}|null – active LLM provider information.\n'
        'currentDate: string – ISO date string used in analysis prompts.'
    ),
    operations=(
        'handleLogin(token): void – saves session token and sets authenticated = true.\n'
        'handleLogout(): void – calls POST /api/logout, clears token, sets authenticated = false.\n'
        'handleUrlQuery(url): void – builds structured change-analysis prompt and calls sendMessage.\n'
        'goHome(): void – clears showBatch and showReports flags.'
    )
)

module_spec(doc,
    name='useChat (Custom React Hook)',
    class_type='Control – chat state manager and SSE client.',
    responsibilities=(
        'Encapsulates the chat messages array and the SSE streaming protocol for /api/chat. '
        'Parses incoming SSE frames and incrementally updates the messages state to drive '
        'real-time rendering.'
    ),
    attributes=(
        'messages: Message[] – array of {role, content, timestamp, tools?} objects, '
        'persisted to localStorage.\n'
        'isLoading: boolean – true while an agentic loop is in progress.\n'
        'activeTools: ToolCall[] – tool calls active in the current response.'
    ),
    operations=(
        'sendMessage(content): Promise<void> – appends user message, POSTs to /api/chat, '
        'and streams SSE events into the messages state.\n'
        'clearChat(): void – resets messages array and clears localStorage.'
    )
)

module_spec(doc,
    name='authFetch (Utility)',
    class_type='Boundary – auth-aware HTTP client.',
    responsibilities=(
        'Wraps the browser fetch API to automatically attach the stored Bearer token '
        'to outgoing requests and invoke the auth-failure handler on 401 responses.'
    ),
    attributes=(
        'onAuthFailure: function|null – global callback invoked when a 401 is received.'
    ),
    operations=(
        'authFetch(url, options?): Promise<Response> – fetch wrapper with automatic auth header.\n'
        'setAuthFailureHandler(handler): void – registers the global 401 callback.'
    )
)

module_spec(doc,
    name='storage (Utility)',
    class_type='Entity – browser storage abstraction.',
    responsibilities=(
        'Provides safe read/write/clear operations on localStorage for chat history, '
        'theme preference, and session token. All operations are wrapped in try/catch '
        'to handle private-browsing or storage-full conditions.'
    ),
    attributes='Keys: "website-monitor-messages", "website-monitor-theme", "website-monitor-token".',
    operations=(
        'loadMessages() / saveMessages(msgs) / clearMessages()\n'
        'loadTheme() / saveTheme(theme)\n'
        'loadToken() / saveToken(token) / clearToken()'
    )
)

# ══════════════════════════════════════════════════════════
#  4. GLOBAL DATA STRUCTURES
# ══════════════════════════════════════════════════════════
add_heading(doc, '4.   GLOBAL DATA STRUCTURES (CSCI DATA)')
add_body(doc, 'The following global data structures are shared across modules:')

gds_headers = ['Name', 'Type', 'Description / Valid Values']
gds_rows = [
    ['Message', 'Object', '{ role: "user"|"assistant", content: string, timestamp: ISO-8601 string, tools?: ToolCall[] }. Used in both API payloads and localStorage persistence.'],
    ['ToolCall', 'Object', '{ name: string, input: object, status: "running"|"done"|"failed", result: string|null }. Attached to assistant messages.'],
    ['SearchResult', 'Object', '{ title: string, url: string, snippet: string }. Returned by webSearchService and serialised as a JSON string for the LLM tool result.'],
    ['AnthropicResponse', 'Object (canonical)', '{ content: ContentBlock[], stop_reason: "end_turn"|"tool_use" }. ContentBlock is { type: "text", text } or { type: "tool_use", id, name, input }.'],
    ['SessionToken', 'string (UUID v4)', 'Randomly generated per login. Stored in activeSessions Set server-side and in localStorage client-side. Valid until server restart or explicit logout.'],
    ['ReportMetadata', 'Object', '{ filename: string, url: string, timestamp: ISO-8601 string, size: number (bytes) }. Returned by GET /api/reports.'],
    ['ENV Config', 'Environment variables', 'LLM_PROVIDER, MODEL_NAME, ANTHROPIC_API_KEY, DEEPSEEK_API_KEY, SEARCH_PROVIDER, BRAVE_SEARCH_API_KEY, SERPAPI_API_KEY, AUTH_USERNAME, AUTH_PASSWORD, MAX_TOOL_CALLS_PER_TURN, MAX_TOKENS, TEMPERATURE, FETCH_MAX_CHARS, WEBSITE_FILE, REPORTS_DIR, PORT, HOST. See .env.example for ranges/defaults.'],
]
add_table_with_header(doc, gds_headers, gds_rows, col_widths=[1.8, 1.2, 3.5])

# ══════════════════════════════════════════════════════════
#  5. CSCI DATA FILES
# ══════════════════════════════════════════════════════════
add_heading(doc, '5.   CSCI DATA FILES')

add_subheading(doc, '5.1   website.md – URL Monitor List')
add_body(doc, 'Plain-text Markdown file; one URL per line. Lines beginning with # are treated as comments. Managed by the WebsiteListEditor via PUT/POST/DELETE /api/batch-monitor/websites.')
p = doc.add_paragraph()
run = p.add_run(
    '# Websites to monitor\n'
    'https://example.com\n'
    'https://news.ycombinator.com\n'
)
p.style = doc.styles['Normal']
run.font.name = 'Courier New'
run.font.size = Pt(10)

add_subheading(doc, '5.2   Report Files (<domain>-<date>.md)')
add_body(doc,
    'Markdown files stored in REPORTS_DIR (default: ./reports/). Each file has a YAML '
    'front-matter header followed by the LLM-generated change report in Markdown.'
)
p = doc.add_paragraph()
run = p.add_run(
    '---\n'
    'url: https://example.com\n'
    'date: 2026-03-06T00:00:00.000Z\n'
    'status: success\n'
    '---\n\n'
    '# Website Change Report: example.com\n\n'
    '<LLM-generated analysis>\n'
)
p.style = doc.styles['Normal']
run.font.name = 'Courier New'
run.font.size = Pt(10)

add_subheading(doc, '5.3   .env Configuration File')
add_body(doc,
    'Environment variable file (not committed to source control). Defines API keys, '
    'provider selection, file paths, and runtime limits. See .env.example for the '
    'full list of variables and their default values.'
)

# ══════════════════════════════════════════════════════════
#  6. REQUIREMENTS TRACEABILITY
# ══════════════════════════════════════════════════════════
add_heading(doc, '6.   REQUIREMENTS TRACEABILITY')
add_body(doc,
    'The table below traces functional requirements to the software components that '
    'implement them. Requirements are derived from the application feature list documented '
    'in README.md.'
)
rt_headers = ['Requirement', 'Source (README Feature)', 'Implementing Component(s)']
rt_rows = [
    ['REQ-01: Single URL change query', 'Single URL Query', 'UrlQueryBar, App.handleUrlQuery, useChat, index.js POST /api/chat, runAgenticLoop, llmService, webSearchService, fetchUrlService'],
    ['REQ-02: Streaming LLM response', 'Streamed analysis', 'index.js SSE (sseWrite), useChat.sendMessageSSE, MessageBubble'],
    ['REQ-03: Batch monitor', 'Batch Monitor', 'BatchMonitorPanel, index.js POST /api/batch-monitor, runAgenticLoop'],
    ['REQ-04: URL list management', 'Website List Editor', 'WebsiteListEditor, index.js GET/PUT/POST/DELETE /api/batch-monitor/websites, readWebsiteFile, writeWebsiteFile'],
    ['REQ-05: Report persistence', 'Reports', 'index.js writeReport, generateReportFilename, GET/DELETE /api/reports'],
    ['REQ-06: Report viewing', 'Reports', 'ReportsPanel, GET /api/reports/:filename'],
    ['REQ-07: Free-form chat', 'Chat', 'ChatInput, useChat, index.js POST /api/chat'],
    ['REQ-08: Dark mode', 'Dark Mode', 'App (darkMode state), Header, Tailwind CSS dark: classes, storage.saveTheme'],
    ['REQ-09: Authentication / login', 'Authentication', 'LoginPage, App, index.js POST /api/login, requireAuth middleware, activeSessions'],
    ['REQ-10: Change password', 'Authentication', 'ChangePasswordModal, index.js POST /api/auth/change-password'],
    ['REQ-11: PDF export', 'Export PDF', 'ReportsPanel (window.print()), browser print dialog'],
    ['REQ-12: LLM provider configurability', 'LLM Provider', 'llmService (LLM_PROVIDER env var, Anthropic/DeepSeek providers)'],
    ['REQ-13: Search provider configurability', 'Search Provider', 'webSearchService (SEARCH_PROVIDER env var, Brave/SerpAPI)'],
    ['REQ-14: SSRF protection', 'Security (implicit)', 'fetchUrlService.isBlockedUrl, validateUrl in index.js'],
    ['REQ-15: Render.com deployment', 'Deploying to Render.com', 'render.yaml, package.json (root scripts), static serving in index.js'],
]
add_table_with_header(doc, rt_headers, rt_rows, col_widths=[1.5, 1.7, 3.3])

# ══════════════════════════════════════════════════════════
#  7. SPECIAL NOTES
# ══════════════════════════════════════════════════════════
add_heading(doc, '7.   SPECIAL NOTES')

add_subheading(doc, '7.1   Agentic Loop Design Decision')
add_body(doc,
    'The agentic loop in runAgenticLoop is bounded by MAX_TOOL_CALLS (default 10) to prevent '
    'runaway tool-calling chains. When the limit is exceeded the loop returns an error tool '
    'result and allows the LLM to conclude its response naturally. This balances thoroughness '
    'of analysis with cost and latency constraints.'
)

add_subheading(doc, '7.2   Provider Normalisation')
add_body(doc,
    'DeepSeek exposes an OpenAI-compatible API. The llmService converts the Anthropic-format '
    'message history used internally by the agentic loop into OpenAI format before calling '
    'DeepSeek, and converts the OpenAI-format response back to Anthropic format on return. '
    'This ensures the agentic loop and all higher-level code remains provider-agnostic.'
)

add_subheading(doc, '7.3   Session Management Limitation')
add_body(doc,
    'Session tokens are stored in a server-side in-memory Set. All sessions are invalidated '
    'on server restart or redeployment. This is an accepted limitation for the initial release; '
    'users are expected to log in again after a server restart. The free Render.com tier spins '
    'down services after a period of inactivity (~15 minutes), so re-authentication will be '
    'required after spin-up on the first request following an idle period.'
)

add_subheading(doc, '7.4   Content Truncation and LLM Context Window')
add_body(doc,
    'Fetched page content is truncated to FETCH_MAX_CHARS (default 8 000 characters). This '
    'prevents exceeding the LLM\'s effective context window when multiple tool results are '
    'accumulated across iterations. Operators may increase this limit for larger-context models '
    'by setting the FETCH_MAX_CHARS environment variable.'
)

add_subheading(doc, '7.5   Security Boundary – SSRF')
add_body(doc,
    'The fetchUrlService and the URL validation function in index.js enforce the following '
    'constraints to prevent Server-Side Request Forgery (SSRF):\n'
    '  • Only HTTPS URLs are accepted.\n'
    '  • Hostnames matching RFC-1918 private ranges (10.x, 172.16–31.x, 192.168.x), '
    'loopback (127.x, ::1), link-local (169.254.x, fe80::), and IPv6 unique-local (fc/fd) '
    'are blocked.\n'
    'This boundary is enforced both at request validation time (validateUrl, before saving '
    'to website.md) and at fetch time (isBlockedUrl, before making the outbound HTTP request).'
)

add_subheading(doc, '7.6   Deployment Persistence')
add_body(doc,
    'On Render.com, the website.md URL list and the reports/ directory are stored on a '
    'persistent 1 GB disk mounted at /data. The WEBSITE_FILE and REPORTS_DIR environment '
    'variables are configured to point to /data/website.md and /data/reports respectively, '
    'ensuring data survives redeployments.'
)

# Save
out_path = '/home/user/website_query4/SDD.docx'
doc.save(out_path)
print(f'Saved: {out_path}')
